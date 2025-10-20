

# CARREGAMENTO DE PACOTES E MODULOS 
import pandas as pd
import numpy as np
import os
from docx import Document
from docx2pdf import convert
from PyPDF2 import PdfReader, PdfWriter
import re
import xlsxwriter
# === Caminho base do script ===
base_dir = os.path.dirname(os.path.abspath(__file__))

# === LENDO A BASE DE DADOS ===
data_frame = pd.read_excel(os.path.join(base_dir, 'BASE 2025.xlsx'))


# === FILTRANDO COLUNAS NECESSÁRIAS ===
filtro_colunas = ['Nome', 'Data de nascimento', 'CPF', 'RG', 'Orgão expedidor',    
       'Data de expedição', 'PIS/NIS', 'E-mail Corporativo', 'Telefone',
       'Sigla', 'Perfil Contrato ']

data_frame_2 = data_frame[filtro_colunas].copy()

# --- NOVO TRATAMENTO ADICIONADO ---
# Converte as colunas de data para o tipo datetime.
# errors='coerce' transforma dados inválidos em NaT (Not a Time) sem gerar erro.
data_frame_2['Data de nascimento'] = pd.to_datetime(data_frame_2['Data de nascimento'], errors='coerce').dt.date
data_frame_2['Data de expedição'] = pd.to_datetime(data_frame_2['Data de expedição'], errors='coerce').dt.date


# === APLICANDO FILTRO DE PERFIL ===
data_frame_2['Perfil Contrato '] = data_frame_2['Perfil Contrato '].str.strip().str.lower()
perfil_desejado = 'analista/desenvolvedor - alta plataforma'
data_frame_filtrado_final = data_frame_2[data_frame_2['Perfil Contrato '] == perfil_desejado]

# # === SALVANDO PLANILHA FILTRADA ===
nome_arquivo_saida = 'CARGOS_FILTRADOS.xlsx'

# UTILIZANDO O EXCELWRITER PARA GARANTIR O FORMATO DATA NO ARQUIVO DE SAIDA
with pd.ExcelWriter(os.path.join(base_dir, nome_arquivo_saida),
                    engine='xlsxwriter',
                    date_format='dd/mm/yyyy') as writer:
    data_frame_filtrado_final.to_excel(writer, index=False)


# === Função de substituição de texto no Word ===
def substituir_texto(doc, antigo, novo):
    def replace_in_runs(runs, antigo, novo):
        for run in runs:
            if antigo in run.text:
                run.text = run.text.replace(antigo, str(novo))
    for p in doc.paragraphs:
        replace_in_runs(p.runs, antigo, novo)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for p in celula.paragraphs:
                    replace_in_runs(p.runs, antigo, novo)

# === Limpeza de nome para evitar erros com caracteres inválidos ===
def limpar_nome(nome):
    return re.sub(r'[\\/*?:"<>|]', "_", nome)

# === LENDO OS DADOS DO ARQUIVO FILTRADO ===
df = pd.read_excel(os.path.join(base_dir, nome_arquivo_saida))
modelo_path = os.path.join(base_dir, "FICUS.docx")

# === Limpa o nome do perfil para usar como nome de diretório ===
perfil_dir_limpo = limpar_nome(perfil_desejado)

# === LOOP PARA CADA COLABORADOR ===
for idx, dados in df.iterrows():
    try:
        colaborador = dados["Nome"].strip()
        colaborador_limpo = limpar_nome(colaborador)
        print(f"📄 Gerando documento para: {colaborador}")

        # === Criar estrutura de pastas: /NOME_DO_PERFIL/NOME_DO_COLABORADOR ===
        pasta_colab = os.path.join(base_dir, perfil_dir_limpo, colaborador_limpo)
        os.makedirs(pasta_colab, exist_ok=True)

        # === Abrir modelo e substituir dados ===
        doc = Document(modelo_path)

        # --- Preenchimento dos dados do colaborador ---
        substituir_texto(doc, "COLABORADOR", colaborador)
        substituir_texto(doc, "000.000.000-00", dados["CPF"])
        substituir_texto(doc, "000000000", dados["RG"])
        substituir_texto(doc, "SSP SP", dados["Orgão expedidor"])
        substituir_texto(doc, "ANALISTA DEV POWER BUILDER", perfil_desejado.upper())
        substituir_texto(doc, "000.00000.00.0", dados["PIS/NIS"])
        substituir_texto(doc, "fulano@gmail.com", dados["E-mail Corporativo"])
        substituir_texto(doc, "000-000-00", dados["Telefone"])
        substituir_texto(doc, "ABC", dados["Sigla"])

        # --- Preenchimento de datas por componentes (dia, mês, ano) ---
        # Data de Nascimento
        if pd.notna(dados["Data de nascimento"]):
            substituir_texto(doc, "dia_nasc", dados["Data de nascimento"].strftime('%d'))
            substituir_texto(doc, "mes_nasc", dados["Data de nascimento"].strftime('%m'))
            substituir_texto(doc, "ano_nasc", dados["Data de nascimento"].strftime('%Y'))
        else:
            substituir_texto(doc, "dia_nasc", "")
            substituir_texto(doc, "mes_nasc", "")
            substituir_texto(doc, "ano_nasc", "")

        # Data de Expedição
        if pd.notna(dados["Data de expedição"]):
            substituir_texto(doc, "dia_expedicao", dados["Data de expedição"].strftime('%d'))
            substituir_texto(doc, "mes_expedicao", dados["Data de expedição"].strftime('%m'))
            substituir_texto(doc, "ano_expedicao", dados["Data de expedição"].strftime('%Y'))
        else:
            substituir_texto(doc, "dia_expedicao", "")
            substituir_texto(doc, "mes_expedicao", "")
            substituir_texto(doc, "ano_expedicao", "")

        # === Salvar DOCX personalizado ===
        novo_docx = os.path.join(pasta_colab, f"{colaborador_limpo}_FICUS.docx")
        doc.save(novo_docx)

        # === Converter para PDF com docx2pdf ===
        novo_pdf = os.path.join(pasta_colab, f"{colaborador_limpo}_FICUS.pdf")
        convert(novo_docx, novo_pdf)

    except Exception as e:
        print(f"❌ Erro ao gerar para {colaborador}: {e}")

print("✅ Todos os PDFs foram gerados com sucesso!")
